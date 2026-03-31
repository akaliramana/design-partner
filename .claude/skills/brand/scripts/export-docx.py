#!/usr/bin/env python3
"""
Export branded DOCX template from MASTER.md tokens using python-docx.

Creates a document with custom heading/body styles matching the project's
design system. HTML is the design source of truth; DOCX is for sharing.

Usage:
    python3 export-docx.py --master design-system/MASTER.md --title "Document Title" --output output.docx

Dependencies:
    pip install python-docx
"""

import argparse
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def hex_to_rgb(hex_color: str) -> tuple:
    """Convert #RRGGBB to (R, G, B) tuple."""
    hex_color = hex_color.lstrip("#")
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def parse_master_tokens(master_path: str) -> dict:
    """Extract brand tokens from MASTER.md."""
    content = Path(master_path).read_text(encoding="utf-8")
    tokens = {}

    for pattern, key in [
        (r"Primary[^\n]*`(#[0-9A-Fa-f]{6})`", "primary"),
        (r"Text[^\n]*`(#[0-9A-Fa-f]{6})`", "text"),
        (r"Background[^\n]*`(#[0-9A-Fa-f]{6})`", "background"),
    ]:
        match = re.search(pattern, content, re.IGNORECASE)
        if match:
            tokens[key] = match.group(1)

    heading = re.search(r"Heading\s*(?:Font)?[:\s]*\**\s*([A-Za-z\s]+)", content)
    body = re.search(r"Body\s*(?:Font)?[:\s]*\**\s*([A-Za-z\s]+)", content)
    tokens["font_heading"] = heading.group(1).strip().rstrip("*") if heading else "Inter"
    tokens["font_body"] = body.group(1).strip().rstrip("*") if body else "Inter"

    return tokens


def create_branded_docx(tokens: dict, title: str, output_path: str):
    """Create a branded DOCX with custom styles from design tokens."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    primary = hex_to_rgb(tokens.get("primary", "#2563EB"))
    text_color = hex_to_rgb(tokens.get("text", "#1E293B"))

    # Style: Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = tokens["font_heading"]
    h1.font.size = Pt(24)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(*text_color)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(8)

    # Style: Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = tokens["font_heading"]
    h2.font.size = Pt(18)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(*text_color)
    h2.paragraph_format.space_before = Pt(20)
    h2.paragraph_format.space_after = Pt(6)

    # Style: Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.name = tokens["font_heading"]
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(*text_color)
    h3.paragraph_format.space_before = Pt(16)
    h3.paragraph_format.space_after = Pt(4)

    # Style: Normal (body)
    normal = doc.styles["Normal"]
    normal.font.name = tokens["font_body"]
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(*text_color)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    # Document title
    title_para = doc.add_heading(title, level=0)
    title_para.runs[0].font.size = Pt(28)
    title_para.runs[0].font.name = tokens["font_heading"]
    title_para.runs[0].font.color.rgb = RGBColor(*text_color)

    # Accent line
    accent_para = doc.add_paragraph()
    accent_run = accent_para.add_run("_" * 20)
    accent_run.font.color.rgb = RGBColor(*primary)
    accent_run.font.size = Pt(4)

    # Placeholder content
    doc.add_paragraph("")
    doc.add_heading("Section Title", level=1)
    doc.add_paragraph("Replace this with your content. All heading and body styles are pre-configured with your brand fonts and colors.")
    doc.add_heading("Subsection", level=2)
    doc.add_paragraph("Body text uses your brand's body font at 11pt with 1.5 line spacing.")

    # Save
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    print(f"Created {output_path}")
    print(f"  Heading font: {tokens['font_heading']}")
    print(f"  Body font: {tokens['font_body']}")
    print(f"  Primary color: {tokens.get('primary', '#2563EB')}")
    print(f"\nOpen in Word and replace placeholder content. Brand styles are pre-applied.")


def main():
    if not HAS_DOCX:
        print("Error: python-docx not installed.")
        print("Install with: pip install python-docx")
        return

    parser = argparse.ArgumentParser(description="Export branded DOCX from MASTER.md tokens")
    parser.add_argument("--master", required=True, help="Path to design-system/MASTER.md")
    parser.add_argument("--title", default="Untitled Document", help="Document title")
    parser.add_argument("--output", default="templates/doc-master.docx", help="Output DOCX path")
    args = parser.parse_args()

    if not Path(args.master).exists():
        print(f"Error: {args.master} not found.")
        return

    tokens = parse_master_tokens(args.master)
    create_branded_docx(tokens, args.title, args.output)


if __name__ == "__main__":
    main()
